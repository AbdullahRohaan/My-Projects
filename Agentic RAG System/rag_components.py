"""
RAG Components - Document Processing, Embedding, and Retrieval
Handles the traditional RAG pipeline components.
"""

import os
import hashlib
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import chromadb
from chromadb.config import Settings
from sentence_transformers import SentenceTransformer
import numpy as np
from pathlib import Path
import pypdf
import docx
import openpyxl
import markdown
from bs4 import BeautifulSoup


@dataclass
class Document:
    """Represents a document in the system"""
    content: str
    metadata: Dict[str, Any]
    doc_id: str
    embedding: Optional[List[float]] = None


class DocumentProcessor:
    """Processes various document types into text"""
    
    @staticmethod
    def process_pdf(file_path: str) -> List[Document]:
        """Extract text from PDF"""
        documents = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                
                if text.strip():
                    doc_id = hashlib.md5(
                        f"{file_path}_page_{page_num}".encode()
                    ).hexdigest()
                    
                    documents.append(Document(
                        content=text,
                        metadata={
                            "source": file_path,
                            "page": page_num + 1,
                            "type": "pdf"
                        },
                        doc_id=doc_id
                    ))
        
        return documents
    
    @staticmethod
    def process_docx(file_path: str) -> List[Document]:
        """Extract text from DOCX"""
        doc = docx.Document(file_path)
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        content = "\n\n".join(paragraphs)
        doc_id = hashlib.md5(file_path.encode()).hexdigest()
        
        return [Document(
            content=content,
            metadata={
                "source": file_path,
                "type": "docx"
            },
            doc_id=doc_id
        )]
    
    @staticmethod
    def process_txt(file_path: str) -> List[Document]:
        """Extract text from TXT"""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        doc_id = hashlib.md5(file_path.encode()).hexdigest()
        
        return [Document(
            content=content,
            metadata={
                "source": file_path,
                "type": "txt"
            },
            doc_id=doc_id
        )]
    
    @staticmethod
    def process_markdown(file_path: str) -> List[Document]:
        """Extract text from Markdown"""
        with open(file_path, 'r', encoding='utf-8') as file:
            md_content = file.read()
        
        # Convert markdown to HTML then extract text
        html = markdown.markdown(md_content)
        soup = BeautifulSoup(html, 'html.parser')
        content = soup.get_text()
        
        doc_id = hashlib.md5(file_path.encode()).hexdigest()
        
        return [Document(
            content=content,
            metadata={
                "source": file_path,
                "type": "markdown"
            },
            doc_id=doc_id
        )]
    
    @classmethod
    def process_file(cls, file_path: str) -> List[Document]:
        """Process a file based on its extension"""
        extension = Path(file_path).suffix.lower()
        
        processors = {
            '.pdf': cls.process_pdf,
            '.docx': cls.process_docx,
            '.txt': cls.process_txt,
            '.md': cls.process_markdown,
        }
        
        processor = processors.get(extension)
        if processor:
            return processor(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")


class TextChunker:
    """Splits documents into manageable chunks"""
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks"""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk = ' '.join(words[i:i + self.chunk_size])
            if chunk:
                chunks.append(chunk)
        
        return chunks
    
    def chunk_documents(self, documents: List[Document]) -> List[Document]:
        """Chunk multiple documents"""
        chunked_docs = []
        
        for doc in documents:
            chunks = self.chunk_text(doc.content)
            
            for idx, chunk in enumerate(chunks):
                chunk_id = f"{doc.doc_id}_chunk_{idx}"
                
                chunked_docs.append(Document(
                    content=chunk,
                    metadata={
                        **doc.metadata,
                        "chunk_index": idx,
                        "parent_doc_id": doc.doc_id
                    },
                    doc_id=chunk_id
                ))
        
        return chunked_docs


class EmbeddingModel:
    """Handles document embedding generation"""
    
    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        self.model = SentenceTransformer(model_name)
        self.dimension = self.model.get_sentence_embedding_dimension()
    
    def embed_text(self, text: str) -> List[float]:
        """Generate embedding for a single text"""
        embedding = self.model.encode(text, convert_to_numpy=True)
        return embedding.tolist()
    
    def embed_batch(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for multiple texts"""
        embeddings = self.model.encode(texts, convert_to_numpy=True)
        return embeddings.tolist()


class VectorStore:
    """Manages the vector database for document retrieval"""
    
    def __init__(self, persist_directory: str = "./chroma_db"):
        self.persist_directory = persist_directory
        self.client = chromadb.PersistentClient(path=persist_directory)
        self.embedding_model = EmbeddingModel()
        
        # Create or get collection
        self.collection = self.client.get_or_create_collection(
            name="documents",
            metadata={"hnsw:space": "cosine"}
        )
    
    def add_documents(self, documents: List[Document]):
        """Add documents to the vector store"""
        if not documents:
            return
        
        # Generate embeddings
        texts = [doc.content for doc in documents]
        embeddings = self.embedding_model.embed_batch(texts)
        
        # Prepare data for insertion
        ids = [doc.doc_id for doc in documents]
        metadatas = [doc.metadata for doc in documents]
        
        # Add to ChromaDB
        self.collection.add(
            ids=ids,
            embeddings=embeddings,
            documents=texts,
            metadatas=metadatas
        )
    
    def search(
        self,
        query: str,
        top_k: int = 3,
        filter_metadata: Optional[Dict[str, Any]] = None
    ) -> List[Dict[str, Any]]:
        """Search for similar documents"""
        
        # Generate query embedding
        query_embedding = self.embedding_model.embed_text(query)
        
        # Search
        results = self.collection.query(
            query_embeddings=[query_embedding],
            n_results=top_k,
            where=filter_metadata
        )
        
        # Format results
        documents = []
        if results['documents'] and results['documents'][0]:
            for i in range(len(results['documents'][0])):
                documents.append({
                    "content": results['documents'][0][i],
                    "metadata": results['metadatas'][0][i] if results['metadatas'] else {},
                    "score": 1 - results['distances'][0][i] if results['distances'] else 0.0,
                    "id": results['ids'][0][i]
                })
        
        return documents
    
    def delete_all(self):
        """Clear all documents from the collection"""
        self.client.delete_collection("documents")
        self.collection = self.client.create_collection(
            name="documents",
            metadata={"hnsw:space": "cosine"}
        )
    
    def get_stats(self) -> Dict[str, Any]:
        """Get statistics about the vector store"""
        count = self.collection.count()
        return {
            "total_documents": count,
            "collection_name": self.collection.name,
            "embedding_dimension": self.embedding_model.dimension
        }


class RAGPipeline:
    """Orchestrates the RAG process"""
    
    def __init__(self, vector_store: VectorStore, chunker: TextChunker = None):
        self.vector_store = vector_store
        self.chunker = chunker or TextChunker()
        self.processor = DocumentProcessor()
    
    def ingest_file(self, file_path: str) -> int:
        """Ingest a single file into the system"""
        # Process file
        documents = self.processor.process_file(file_path)
        
        # Chunk documents
        chunked_docs = self.chunker.chunk_documents(documents)
        
        # Add to vector store
        self.vector_store.add_documents(chunked_docs)
        
        return len(chunked_docs)
    
    def ingest_directory(self, directory_path: str) -> Dict[str, int]:
        """Ingest all supported files from a directory"""
        stats = {"total_chunks": 0, "files_processed": 0, "errors": 0}
        
        for file_path in Path(directory_path).rglob("*"):
            if file_path.is_file():
                try:
                    chunks = self.ingest_file(str(file_path))
                    stats["total_chunks"] += chunks
                    stats["files_processed"] += 1
                except Exception as e:
                    print(f"Error processing {file_path}: {e}")
                    stats["errors"] += 1
        
        return stats
    
    async def retrieve(self, query: str, top_k: int = 3) -> List[Dict[str, Any]]:
        """Retrieve relevant documents for a query"""
        return self.vector_store.search(query, top_k)
